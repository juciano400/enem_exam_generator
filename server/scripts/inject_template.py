#!/usr/bin/env python3
"""
inject_template.py — Lê um template .docx, entende sua estrutura e injeta
questões ENEM no ponto correto, preservando o visual do layout original.

Marcadores suportados no template (coloque no local onde as questões devem entrar):
  {{QUESTOES}}  ou  {{QUESTIONS}}  ou  [QUESTOES]  ou  QUESTOES_AQUI

Se nenhum marcador for encontrado, as questões são inseridas antes dos
parágrafos vazios finais (fim da área de conteúdo).
"""

import sys, json, os, re, argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Cores ──────────────────────────────────────────────────────────────────────
BLUE_DARK = RGBColor(0x1A, 0x3A, 0x6E)
BLUE_MID  = RGBColor(0x44, 0x72, 0xC4)
GRAY_CTX  = RGBColor(0x55, 0x55, 0x66)
GREEN_ANS = RGBColor(0x0A, 0x7A, 0x2A)
GRAY_SEP  = RGBColor(0xCC, 0xCC, 0xCC)

# Letras circuladas Unicode (Ⓐ Ⓑ Ⓒ Ⓓ Ⓔ)
CIRCLED = {'A': 'Ⓐ', 'B': 'Ⓑ', 'C': 'Ⓒ', 'D': 'Ⓓ', 'E': 'Ⓔ'}


# ── Detecção de fonte base ─────────────────────────────────────────────────────

def detect_base_font(doc):
    """Retorna (nome_fonte, tamanho_pt) do documento."""
    name, size = 'Arial', 11.0
    try:
        s = doc.styles['Normal']
        if s.font.name:  name = s.font.name
        if s.font.size:  size = s.font.size.pt
    except Exception:
        pass
    for para in doc.paragraphs[:15]:
        for run in para.runs:
            if run.font.name and run.font.size:
                name = run.font.name
                size = run.font.size.pt
                return name, size
    return name, size


# ── Preenchimento de campos do cabeçalho ──────────────────────────────────────

FIELD_PATTERNS = [
    # (regex a detectar, texto de substituição com placeholder {v})
    (r'COLOQUE\s+A\s+SUA\s+DISCIPLINA|DISCIPLINA\s*:\s*$|^\s*DISCIPLINA\s*$',
     lambda v: v),
    (r'(?:TOTAL|Nº|N[°º]|NUM(?:ERO)?\.?)\s+DE\s+QUESTÕES?\s*:?\s*\d*',
     lambda v: f'TOTAL DE QUESTÕES: {v}'),
    (r'QUESTÕES?\s*:\s*\d*\s*$',
     lambda v: f'QUESTÕES: {v}'),
]


def _replace_in_paragraph(para, pattern, replacement_text, font_name):
    """Substitui texto em um parágrafo preservando formatação do primeiro run."""
    full = ''.join(r.text for r in para.runs)
    if not re.search(pattern, full, re.IGNORECASE):
        return False
    # Zera todos os runs e coloca o texto no primeiro
    for r in para.runs:
        r.text = ''
    target = para.runs[0] if para.runs else para.add_run()
    target.text = replacement_text
    if font_name:
        target.font.name = font_name
    return True


def fill_header_fields(doc, discipline, question_count, serie=None, turma=None, gabarito=False):
    """Preenche campos padrão em tabelas e parágrafos do cabeçalho."""
    font_name, _ = detect_base_font(doc)
    disc_label = discipline.upper() + (' — GABARITO' if gabarito else '')

    substitutions = [
        (r'COLOQUE\s+A\s+SUA\s+DISCIPLINA|(?:^|\b)DISCIPLINA\s*:\s*$|^\s*DISCIPLINA\s*$',
         disc_label),
        (r'(?:TOTAL|Nº|N[°º])\s+DE\s+QUESTÕES?\s*:?\s*\d*',
         f'TOTAL DE QUESTÕES: {question_count}'),
        (r'QUESTÕES?\s*:\s*\d*\s*$',
         f'QUESTÕES: {question_count}'),
    ]
    if serie:
        substitutions.append((r'SÉRIE\s*:?\s*[_\s]*$', f'SÉRIE: {serie}'))
    if turma:
        substitutions.append((r'TURMA\s*:?\s*[_\s]*$', f'TURMA: {turma}'))

    # Tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for pattern, replacement in substitutions:
                    if re.search(pattern, text, re.IGNORECASE):
                        for p in cell.paragraphs:
                            _replace_in_paragraph(p, pattern, replacement, font_name)
                        break

    # Parágrafos soltos
    for para in doc.paragraphs:
        for pattern, replacement in substitutions:
            if _replace_in_paragraph(para, pattern, replacement, font_name):
                break


# ── Localização do ponto de inserção ──────────────────────────────────────────

MARKER_RE = re.compile(
    r'\{\{QUESTOES?\}\}|\{QUESTOES?\}|\[QUESTOES?\]|QUESTOES?_AQUI|INSERT_QUESTIONS?',
    re.IGNORECASE
)


def find_marker_paragraph(doc):
    """Retorna o parágrafo marcador ou None."""
    for para in doc.paragraphs:
        if MARKER_RE.search(para.text):
            return para
    return None


def find_fallback_insertion(doc):
    """
    Fallback: retorna o último parágrafo não-vazio do documento.
    As questões serão inseridas APÓS esse parágrafo.
    """
    last = None
    for para in doc.paragraphs:
        if para.text.strip():
            last = para
    return last


# ── Construção de parágrafos OOXML ────────────────────────────────────────────

def _new_p():
    return OxmlElement('w:p')


def _pPr(**attrs):
    """Cria um elemento pPr com atributos filhos."""
    pPr = OxmlElement('w:pPr')
    if 'spacing' in attrs:
        spc = OxmlElement('w:spacing')
        for k, v in attrs['spacing'].items():
            spc.set(qn(f'w:{k}'), str(v))
        pPr.append(spc)
    if 'ind' in attrs:
        ind = OxmlElement('w:ind')
        for k, v in attrs['ind'].items():
            ind.set(qn(f'w:{k}'), str(v))
        pPr.append(ind)
    if 'borders' in attrs:
        pBdr = OxmlElement('w:pBdr')
        for side, props in attrs['borders'].items():
            el = OxmlElement(f'w:{side}')
            for k, v in props.items():
                el.set(qn(f'w:{k}'), str(v))
            pBdr.append(el)
        pPr.append(pBdr)
    return pPr


def _add_run(p_el, text, bold=False, italic=False, size_pt=11,
             font_name='Arial', color: RGBColor = None):
    """Adiciona um run a um elemento <w:p>."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Fonte
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    # Tamanho (em half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szCs)

    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    if color:
        clr = OxmlElement('w:color')
        clr.set(qn('w:val'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
        rPr.append(clr)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p_el.append(r)
    return r


def _insert_after(anchor_el, new_el):
    """Insere new_el imediatamente após anchor_el no mesmo pai."""
    anchor_el.addnext(new_el)


# ── Injeção de questões ────────────────────────────────────────────────────────

def inject_questions(doc, anchor_para, questions, font_name, font_size, gabarito=False):
    """
    Injeta questões usando o anchor_para como referência.
    Se anchor_para é um marcador → substitui-o.
    Caso contrário → insere após ele.
    """
    anchor_p = anchor_para._p
    is_marker = bool(MARKER_RE.search(anchor_para.text))

    # Vamos inserir sempre APÓS o anchor (ou no lugar do marker após removê-lo)
    # Usamos uma "cauda" que vai sendo atualizada a cada inserção.
    if is_marker:
        # Inserir no lugar do marcador: primeiro inserimos algo antes, depois removemos
        placeholder = _new_p()
        anchor_p.addprevious(placeholder)
        current = placeholder          # inserir "após" o placeholder
        anchor_p.getparent().remove(anchor_p)
    else:
        current = anchor_p            # inserir após o último parágrafo de conteúdo

    def emit(p_el):
        nonlocal current
        _insert_after(current, p_el)
        current = p_el

    for q in questions:
        num       = q.get('number', '?')
        context   = (q.get('context') or '').strip()
        statement = (q.get('statement') or '').strip()
        alts      = q.get('alternatives', [])
        correct   = q.get('correctAnswer', '?')
        explanation = (q.get('explanation') or '').strip()

        # ── Cabeçalho: "QUESTÃO N" ──────────────────────────────────────
        p_hdr = _new_p()
        p_hdr.append(_pPr(spacing={'before': '100', 'after': '30'}))
        _add_run(p_hdr, f'QUESTÃO {num}',
                 bold=True, size_pt=font_size, font_name=font_name, color=BLUE_DARK)
        if gabarito:
            _add_run(p_hdr, f'   →  Resposta: {correct}',
                     bold=True, size_pt=font_size, font_name=font_name, color=GREEN_ANS)
        emit(p_hdr)

        # Linha separadora azul sob o cabeçalho
        p_line = _new_p()
        p_line.append(_pPr(
            spacing={'before': '0', 'after': '40'},
            borders={'bottom': {'val': 'single', 'sz': '6', 'space': '1', 'color': '1A3A6E'}}
        ))
        emit(p_line)

        # ── Contexto (texto de apoio) ────────────────────────────────────
        if context and not gabarito:
            p_ctx = _new_p()
            p_ctx.append(_pPr(
                spacing={'before': '40', 'after': '60'},
                ind={'left': '360'},
                borders={'left': {'val': 'single', 'sz': '12', 'space': '4', 'color': '4472C4'}}
            ))
            _add_run(p_ctx, context, italic=True,
                     size_pt=max(font_size - 1, 9), font_name=font_name, color=GRAY_CTX)
            emit(p_ctx)

        # ── Enunciado ────────────────────────────────────────────────────
        p_stmt = _new_p()
        p_stmt.append(_pPr(spacing={'before': '60', 'after': '60'}))
        if gabarito:
            short = statement[:130] + ('…' if len(statement) > 130 else '')
            _add_run(p_stmt, short, italic=True,
                     size_pt=font_size - 0.5, font_name=font_name, color=GRAY_CTX)
        else:
            _add_run(p_stmt, statement, bold=True,
                     size_pt=font_size, font_name=font_name)
        emit(p_stmt)

        if gabarito:
            # Explicação
            p_lbl = _new_p()
            p_lbl.append(_pPr(spacing={'before': '20', 'after': '20'}, ind={'left': '280'}))
            _add_run(p_lbl, 'Explicação: ', bold=True,
                     size_pt=font_size - 0.5, font_name=font_name)
            emit(p_lbl)

            p_exp = _new_p()
            p_exp.append(_pPr(spacing={'before': '0', 'after': '60'}, ind={'left': '360'}))
            _add_run(p_exp, explanation,
                     size_pt=font_size - 0.5, font_name=font_name)
            emit(p_exp)
        else:
            # ── Alternativas ─────────────────────────────────────────────
            for alt in alts:
                letter = alt.get('letter', '?')
                text   = (alt.get('text') or '').strip()
                circle = CIRCLED.get(letter, f'({letter})')

                p_alt = _new_p()
                p_alt.append(_pPr(
                    spacing={'before': '20', 'after': '20'},
                    ind={'left': '440', 'hanging': '440'}
                ))
                _add_run(p_alt, f'{circle}  ', bold=True,
                         size_pt=font_size, font_name=font_name, color=BLUE_DARK)
                _add_run(p_alt, text,
                         size_pt=font_size, font_name=font_name)
                emit(p_alt)

        # ── Separador final ───────────────────────────────────────────────
        p_sep = _new_p()
        p_sep.append(_pPr(
            spacing={'before': '60', 'after': '60'},
            borders={'bottom': {'val': 'single', 'sz': '4', 'space': '1', 'color': 'CCCCCC'}}
        ))
        emit(p_sep)

    # Remove o placeholder vazio se foi criado
    if is_marker:
        placeholder.getparent().remove(placeholder)


# ── Gabarito resumido (grade compacta no topo) ────────────────────────────────

def add_answer_summary(doc, questions, anchor_para, font_name, font_size):
    """Adiciona um bloco de resumo do gabarito antes das questões detalhadas."""
    anchor_p = anchor_para._p
    current  = anchor_p

    def emit(p_el):
        nonlocal current
        _insert_after(current, p_el)
        current = p_el

    # Título
    p_title = _new_p()
    p_title.append(_pPr(spacing={'before': '80', 'after': '40'}))
    _add_run(p_title, 'GABARITO E EXPLICAÇÕES', bold=True,
             size_pt=font_size + 2, font_name=font_name, color=BLUE_DARK)
    emit(p_title)

    # Linha resumo: "1:A  2:B  3:C ..."
    p_sum = _new_p()
    p_sum.append(_pPr(spacing={'before': '20', 'after': '60'}))
    summary = '   '.join(f"{q['number']}:{q['correctAnswer']}" for q in questions)
    _add_run(p_sum, summary, bold=True,
             size_pt=font_size - 0.5, font_name=font_name, color=BLUE_DARK)
    emit(p_sum)

    # Separador
    p_sep = _new_p()
    p_sep.append(_pPr(
        spacing={'before': '20', 'after': '80'},
        borders={'bottom': {'val': 'single', 'sz': '6', 'space': '1', 'color': '1A3A6E'}}
    ))
    emit(p_sep)

    return current  # novo anchor para as questões detalhadas


# ── Entry point ───────────────────────────────────────────────────────────────

def process(template_path, questions_path, output_path, discipline, gabarito=False,
            serie=None, turma=None):
    with open(questions_path, 'r', encoding='utf-8') as f:
        questions = json.load(f)

    doc = Document(template_path)
    font_name, font_size = detect_base_font(doc)
    print(f'[INFO] Fonte detectada: {font_name} {font_size}pt', file=sys.stderr)

    # 1. Preenche campos do cabeçalho
    fill_header_fields(doc, discipline, len(questions), serie=serie, turma=turma, gabarito=gabarito)

    # 2. Localiza ponto de inserção
    marker = find_marker_paragraph(doc)
    if marker:
        print(f'[INFO] Marcador encontrado: "{marker.text.strip()}"', file=sys.stderr)
        anchor = marker
    else:
        anchor = find_fallback_insertion(doc)
        if anchor:
            print(f'[INFO] Inserindo após último parágrafo de conteúdo.', file=sys.stderr)
        else:
            print('[WARN] Nenhum ponto de inserção encontrado; appendando ao final.', file=sys.stderr)
            anchor = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()

    # 3. Para gabarito: adiciona sumário antes das explicações
    if gabarito:
        summary_anchor = anchor
        inject_anchor = add_answer_summary(doc, questions, summary_anchor, font_name, font_size)
        # Inject detailed explanations after the summary
        # Create a dummy paragraph as anchor for injection
        dummy_p = _new_p()
        _insert_after(inject_anchor, dummy_p)
        from docx.text.paragraph import Paragraph
        dummy_para = Paragraph(dummy_p, doc.element.body)
        inject_questions(doc, dummy_para, questions, font_name, font_size, gabarito=True)
    else:
        inject_questions(doc, anchor, questions, font_name, font_size, gabarito=False)

    doc.save(output_path)
    print(f'[OK] Salvo: {output_path}', file=sys.stderr)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('template')
    parser.add_argument('questions_json')
    parser.add_argument('output_docx')
    parser.add_argument('--discipline', default='')
    parser.add_argument('--gabarito', action='store_true')
    parser.add_argument('--serie', default='')
    parser.add_argument('--turma', default='')
    args = parser.parse_args()

    process(
        args.template, args.questions_json, args.output_docx,
        discipline=args.discipline,
        gabarito=args.gabarito,
        serie=args.serie or None,
        turma=args.turma or None,
    )
    print(args.output_docx)
